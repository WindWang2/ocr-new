import os
import glob
import json
import docx

def main():
    docs_dir = os.path.join("docs", "exp-templates")
    files = glob.glob(os.path.join(docs_dir, "*.docx"))
    
    results = {}
    for f in files:
        try:
            doc = docx.Document(f)
            text_blocks = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_blocks.append(para.text.strip())
                    
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_data:
                        text_blocks.append(" | ".join(row_data))
                        
            results[os.path.basename(f)] = '\n'.join(text_blocks)
        except Exception as e:
            results[os.path.basename(f)] = f"Error: {e}"
            
    with open(os.path.join("scratch", "docs_content_tables.json"), "w", encoding="utf-8") as f:
        json.dump(results, f, ensure_ascii=False, indent=2)

if __name__ == "__main__":
    main()
